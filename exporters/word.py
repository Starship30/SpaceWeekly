from datetime import datetime
from pathlib import Path

from config import OUTPUT_DIR
from i18n import tr
from models.ai_summary import AISummary
from models.article import Article
from models.export_context import ExportContext
from models.news_analysis import NewsAnalysis

IMPORTANCE_LEVELS = ["High", "Medium", "Low"]


def export_word(
    articles: list[Article],
    ai_summaries: dict[str, AISummary] | None = None,
    context: ExportContext | None = None,
    analyses: dict[str, NewsAnalysis] | None = None,
    launches: list[Article] | None = None,
    launch_range: str = "next_week",
    launch_term_translator=None,
) -> Path:
    """Export articles to a styled Word document."""
    from docx import Document

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = _output_path()
    document = Document()
    _setup_styles(document)
    document.add_heading(_report_title(context), level=0)

    if launch_range != "disabled":
        document.add_heading(_launch_section_title(launch_range), level=1)

        if launches:
            for launch in launches:
                _add_launch(document, launch, launch_term_translator)
        else:
            document.add_paragraph(tr("launch_none"))

    document.add_heading(tr("weekly_space_news"), level=1)

    for section, section_articles in _group_by_importance(
        articles,
        analyses or {},
    ).items():
        document.add_heading(section, level=2)

        for article in section_articles:
            _add_article(
                document,
                article,
                (ai_summaries or {}).get(article.news.url),
                context,
                (analyses or {}).get(article.news.url),
            )

    document.save(output_path)

    return output_path


def _add_launch(
    document,
    launch: Article,
    term_translator,
) -> None:
    fields = _launch_fields(launch)
    fields["name"] = _translated_term(fields["name"], "mission", term_translator)
    fields["rocket"] = _translated_term(fields["rocket"], "rocket", term_translator)
    fields["provider"] = _translated_term(
        fields["provider"],
        "organization",
        term_translator,
    )
    fields["pad"] = _translated_term(fields["pad"], "location", term_translator)
    fields["mission_name"] = _translated_term(
        fields["mission_name"],
        "mission",
        term_translator,
    )
    fields["orbit"] = _translated_term(fields["orbit"], "orbit", term_translator)
    fields["status"] = _translated_term(fields["status"], "status", term_translator)
    fields["mission_description"] = _translated_description(
        fields["mission_description"],
        term_translator,
    )
    title = fields["mission_name"]

    if title == tr("not_available"):
        title = fields["name"]

    document.add_heading(f"🚀 {title}", level=2)
    _add_label(document, tr("launch_time"), fields["time"])
    _add_label(document, tr("rocket"), fields["rocket"])
    _add_label(document, tr("launch_operator"), fields["provider"])
    _add_label(document, tr("launch_site"), fields["pad"])
    _add_label(document, tr("orbit"), fields["orbit"])
    _add_label(document, tr("status"), fields["status"])
    _add_label(document, tr("mission_description"), fields["mission_description"])


def _launch_section_title(launch_range: str) -> str:
    if launch_range == "past_week":
        return tr("launch_section_past")

    return tr("launch_section_next")


def _translated_term(
    value: str,
    term_type: str,
    translator,
) -> str:
    if translator is None or value in {"Not available", tr("not_available")}:
        return value

    return translator.translate_term(value, term_type, allow_ai=False) or value


def _translated_description(value: str, translator) -> str:
    if (
        translator is None
        or value in {"Not available", tr("not_available")}
        or value == "No mission description available."
    ):
        return value

    return translator.translate_description(value) or value


def _launch_fields(launch: Article) -> dict[str, str]:
    fields = {
        "name": launch.news.title,
        "time": launch.news.published,
        "rocket": "",
        "provider": "",
        "pad": "",
        "mission_name": "",
        "mission_description": "",
        "orbit": "",
        "status": "",
    }
    labels = {
        "Launch name:": "name",
        "Launch time:": "time",
        "Rocket:": "rocket",
        "Launch service provider:": "provider",
        "Launch pad:": "pad",
        "Mission name:": "mission_name",
        "Mission description:": "mission_description",
        "Orbit:": "orbit",
        "Status:": "status",
        "Official or video link:": None,
    }
    multiline_field: str | None = None

    for raw_line in launch.body.splitlines():
        line = raw_line.strip()

        if not line:
            continue

        matched = False

        for label, key in labels.items():
            if not line.startswith(label):
                continue

            matched = True
            multiline_field = key if key == "mission_description" else None

            if key is not None:
                value = line[len(label):].strip()

                if value:
                    fields[key] = value

            break

        if matched:
            continue

        if multiline_field:
            existing = fields[multiline_field]
            fields[multiline_field] = f"{existing}\n{line}".strip()

    return {key: value or tr("not_available") for key, value in fields.items()}


def _output_path() -> Path:
    today = datetime.now().date().isoformat()

    return OUTPUT_DIR / f"SpaceWeekly_{today}.docx"


def _setup_styles(document) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    for style_name in ["Normal", "Heading 1", "Heading 2"]:
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    document.styles["Normal"].font.size = Pt(11)
    document.styles["Heading 1"].font.size = Pt(18)
    document.styles["Heading 2"].font.size = Pt(14)
    document.styles["Heading 1"].font.bold = True
    document.styles["Heading 2"].font.bold = True


def _report_title(context: ExportContext | None) -> str:
    if context is None:
        return "SpaceWeekly"

    return context.report_title or "SpaceWeekly"


def _add_article(
    document,
    article: Article,
    ai_summary: AISummary | None,
    context: ExportContext | None,
    analysis: NewsAnalysis | None,
) -> None:
    news = article.news

    if context is None or context.include_title:
        document.add_heading(news.title, level=2)

    if context is None or context.include_published:
        _add_label(document, tr("report_published"), news.published)

    if context is None or context.include_source:
        _add_label(document, tr("report_source"), news.source)

    if analysis is not None and (context is None or context.include_score):
        _add_label(document, tr("report_importance"), _importance_label(_importance_level(analysis)))
        _add_label(document, tr("report_reason"), analysis.reason)

    if analysis is not None and (context is None or context.include_categories):
        _add_label(document, tr("report_category"), _join_values(analysis.categories))

    if ai_summary is None:
        _add_label(document, tr("report_summary"), news.summary)
    else:
        _add_ai_summary(document, ai_summary, context)

    if context is None or context.include_body:
        body_text = _body_text(article, context)

        if body_text:
            document.add_heading(tr("report_body"), level=2)
            document.add_paragraph(body_text)

    if context is None or context.include_link:
        paragraph = document.add_paragraph()
        paragraph.add_run(
            f"{tr('report_link')}{tr('label_separator')}"
        ).bold = True
        _add_hyperlink(paragraph, news.url, news.url)


def _add_ai_summary(document, ai_summary: AISummary, context: ExportContext | None) -> None:
    if context is None or (context.show_summary and context.include_summary):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = 6
        run = paragraph.add_run(
            f"{tr('report_ai_summary')}{tr('label_separator')}{ai_summary.summary}"
        )
        run.bold = True

    if context is None or (context.show_keywords and context.include_keywords):
        _add_label(document, tr("report_keywords"), _join_values(ai_summary.keywords))

    if context is None or (context.show_category and context.include_categories):
        _add_label(document, tr("report_category"), ai_summary.category)

    if context is None or (context.show_importance and context.include_score):
        _add_label(document, tr("report_importance"), _importance_label(ai_summary.importance))


def _body_text(article: Article, context: ExportContext | None) -> str:
    translation = context.translations.get(article.news.url) if context else None

    if context is not None and context.body_mode == "none":
        return ""

    if context is not None and context.body_mode == "translated" and translation:
        return translation

    if context is not None and context.body_mode == "bilingual" and translation:
        return _bilingual_body(article.body, translation)

    if context is not None and not context.include_original:
        return ""

    return article.body


def _bilingual_body(original: str, translation: str) -> str:
    original_lines = [line for line in original.splitlines() if line.strip()]
    translated_lines = [line for line in translation.splitlines() if line.strip()]
    blocks = []

    for index, original_line in enumerate(original_lines):
        blocks.append(original_line)

        if index < len(translated_lines):
            blocks.append(translated_lines[index])

        blocks.append("")

    return "\n".join(blocks).strip()


def _add_label(document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"{label}{tr('label_separator')}")
    run.bold = True
    paragraph.add_run(value)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE

    relationship_id = paragraph.part.relate_to(
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._element.append(hyperlink)


def _group_by_importance(
    articles: list[Article],
    analyses: dict[str, NewsAnalysis],
) -> dict[str, list[Article]]:
    grouped = {_importance_label(level): [] for level in IMPORTANCE_LEVELS}

    for article in articles:
        analysis = analyses.get(article.news.url)
        label = _importance_label(_importance_level(analysis))
        grouped.setdefault(label, []).append(article)

    return {label: items for label, items in grouped.items() if items}


def _importance_level(analysis: NewsAnalysis | None) -> str:
    if analysis is None:
        return "Medium"

    text = (analysis.importance or "").strip().lower()
    aliases = {
        "high": "High",
        "高": "High",
        "medium": "Medium",
        "中": "Medium",
        "low": "Low",
        "低": "Low",
    }

    if text in aliases:
        return aliases[text]

    if analysis.score >= 80:
        return "High"

    if analysis.score >= 50:
        return "Medium"

    return "Low"


def _importance_label(level: str) -> str:
    text = (level or "").strip()
    aliases = {
        "High": tr("importance_high"),
        "Medium": tr("importance_medium"),
        "Low": tr("importance_low"),
        "high": tr("importance_high"),
        "medium": tr("importance_medium"),
        "low": tr("importance_low"),
    }

    if text in aliases:
        return aliases[text]

    return text or tr("importance_medium")


def _join_values(values: list[str]) -> str:
    return tr("list_separator").join(values)
